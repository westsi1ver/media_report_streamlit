from __future__ import annotations

import calendar
import copy
import io
import re
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as _Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

BLUE = RGBColor(0x00, 0x00, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)


@dataclass
class WeekReport:
    month: int
    day: int
    week_no: int
    creator_meeting: list[str] = field(default_factory=list)
    whole_team_meeting: list[str] = field(default_factory=list)
    bulletin: dict[str, str] = field(default_factory=dict)
    external: list[str] = field(default_factory=list)
    internal: list[str] = field(default_factory=list)
    other: list[str] = field(default_factory=list)


@dataclass
class MonthlyData:
    year: int
    month: int
    bulletin_rows: list[dict] = field(default_factory=list)
    whole_team: list[str] = field(default_factory=list)
    internal: list[str] = field(default_factory=list)
    external: list[str] = field(default_factory=list)
    finance_total: str | None = None
    prayers: list[str] = field(default_factory=list)


def _clean_line(line: str) -> str:
    line = line.strip().replace('\u00a0', ' ')
    line = re.sub(r'^[○●]\s*', '', line)
    line = re.sub(r'^→\s*', '', line)
    return line.strip()


def parse_weekly_text(text: str) -> list[WeekReport]:
    """Parse the user's weekly media-team text. No content is invented."""
    header_re = re.compile(r'(?:🎨\s*)?미디어팀\s+(\d{1,2})/(\d{1,2})')
    matches = list(header_re.finditer(text))
    reports: list[WeekReport] = []
    for i, m in enumerate(matches):
        month, day = int(m.group(1)), int(m.group(2))
        block = text[m.end(): matches[i + 1].start() if i + 1 < len(matches) else len(text)]
        week = WeekReport(month=month, day=day, week_no=((day - 1) // 7) + 1)
        section = None
        pending_meeting = None
        for raw in block.splitlines():
            s = raw.strip()
            if not s or re.fullmatch(r'[-—–_=\s]+', s):
                continue
            if s.startswith('●'):
                heading = _clean_line(s)
                if heading.startswith('온라인 주보'):
                    section = 'bulletin'; pending_meeting = None
                elif heading == '미디어팀':
                    section = 'team'; pending_meeting = None
                elif heading.startswith('크리에이터/이미지'):
                    section = 'work'; pending_meeting = None
                else:
                    section = 'other'; pending_meeting = None
                continue
            line = _clean_line(s)
            if not line:
                continue
            if section == 'bulletin':
                # e.g. 전면 : ... / 내지 : ...
                mm = re.match(r'(전면|후면|내지|단추가)\s*[:：]\s*(.+)$', line)
                if mm:
                    week.bulletin[mm.group(1)] = mm.group(2).strip()
            elif section == 'team':
                if '크리에이터 모임' in line:
                    pending_meeting = 'creator'
                elif '전체 팀모임' in line or '전체팀' in line:
                    pending_meeting = 'whole'
                else:
                    if pending_meeting == 'creator':
                        week.creator_meeting.append(line)
                    elif pending_meeting == 'whole':
                        # comma-separated items are kept as one source sentence
                        week.whole_team_meeting.append(line)
                    else:
                        week.other.append(line)
            elif section == 'work':
                mm = re.match(r'\[(외부|자체)\]\s*(.+)', line)
                if mm:
                    (week.external if mm.group(1) == '외부' else week.internal).append(mm.group(2).strip())
                else:
                    week.other.append(line)
            else:
                week.other.append(line)
        reports.append(week)
    return reports


def _unique_tcs(row):
    out = []
    seen = set()
    for cell in row.cells:
        key = id(cell._tc)
        if key not in seen:
            out.append(cell)
            seen.add(key)
    return out


def _clear_cell(cell):
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn('w:tcPr'):
            tc.remove(child)
    p = OxmlElement('w:p')
    tc.append(p)


def _set_cell_lines(cell, lines: Iterable[tuple[str, bool]], font_size: float = 9.0, bold_first: bool = False):
    _clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    first = True
    for text, blue in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(text)
        r.font.size = Pt(font_size)
        r.font.color.rgb = BLUE if blue else BLACK
        if bold_first and text.startswith(('[', '<')):
            r.bold = True


def _clone_row_before(table, source_row_idx: int, before_row_idx: int):
    new_tr = copy.deepcopy(table.rows[source_row_idx]._tr)
    table.rows[before_row_idx]._tr.addprevious(new_tr)
    # recreate table row wrappers by accessing table.rows after insertion
    return table.rows[before_row_idx]


def _week_rows(table):
    result = {}
    for idx, row in enumerate(table.rows):
        txt = row.cells[0].text.strip()
        m = re.fullmatch(r'(\d+)주', txt)
        if m:
            result[int(m.group(1))] = idx
    return result


def _finance_header_idx(table):
    for idx, row in enumerate(table.rows):
        if row.cells[0].text.strip() == '재정' and any(c.text.strip() == '구분' for c in row.cells):
            return idx
    return None


def _prayer_row_idx(table):
    for idx, row in enumerate(table.rows):
        texts = [c.text.strip() for c in _unique_tcs(row)]
        if any('기도' in t for t in texts) and any('제목' in t for t in texts):
            return idx
    return None


def _set_title_month(table, month: int):
    # Top merged cells: left church / center title / right team.
    for c in _unique_tcs(table.rows[0]):
        if '보 고 서' in c.text:
            old = c.text
            new = re.sub(r'\d+월', f'{month}월', old)
            for p in c.paragraphs:
                for r in p.runs:
                    if '보 고 서' in r.text or re.search(r'\d+월', r.text):
                        r.text = re.sub(r'\d+월', f'{month}월', r.text)
            if c.text == old:  # fallback
                c.text = new


def build_monthly_docx(template_path: str | Path, weekly_text: str, year: int, month: int) -> bytes:
    reports = [r for r in parse_weekly_text(weekly_text) if r.month == month]
    doc = Document(str(template_path))
    table = doc.tables[0]
    _set_title_month(table, month)

    # Ensure rows 1..5 exist. Clone the last existing week-row just before finance.
    week_map = _week_rows(table)
    finance_idx = _finance_header_idx(table)
    while max(week_map or {0: 0}) < 5:
        source_idx = week_map[max(week_map)]
        finance_idx = _finance_header_idx(table)
        new_week = max(week_map) + 1
        row = _clone_row_before(table, source_idx, finance_idx)
        _set_cell_lines(row.cells[0], [(f'{new_week}주', False)], font_size=10)
        for c in _unique_tcs(row)[1:]:
            _clear_cell(c)
        week_map = _week_rows(table)

    report_by_week = {r.week_no: r for r in reports}
    week_map = _week_rows(table)
    for w in range(1, 6):
        row = table.rows[week_map[w]]
        unique = _unique_tcs(row)
        # first unique cell is week label; remaining are left and right activity cells.
        left = unique[1] if len(unique) > 1 else row.cells[1]
        right = unique[2] if len(unique) > 2 else row.cells[-1]
        _set_cell_lines(row.cells[0], [(f'{w}주', False)], font_size=10)
        rep = report_by_week.get(w)
        if not rep:
            _clear_cell(left); _clear_cell(right)
            continue
        left_lines: list[tuple[str, bool]] = []
        if rep.creator_meeting:
            left_lines.append(('<크리에이터> 팀모임', False))
            left_lines += [(f'- {x}', False) for x in rep.creator_meeting]
        if rep.whole_team_meeting:
            if left_lines: left_lines.append(('', False))
            left_lines.append(('<전체팀> 팀모임', False))
            left_lines += [(f'- {x}', False) for x in rep.whole_team_meeting]
        if rep.bulletin:
            if left_lines: left_lines.append(('', False))
            left_lines.append(('[주보]', False))
            for key in ('전면', '후면', '내지', '단추가'):
                if key in rep.bulletin:
                    left_lines.append((f'- {key} : {rep.bulletin[key]}', False))
        if rep.other:
            if left_lines: left_lines.append(('', False))
            left_lines.append(('[기타]', True))  # synthesized category is visibly blue
            left_lines += [(f'- {x}', False) for x in rep.other]
        right_lines: list[tuple[str, bool]] = []
        if rep.external:
            right_lines.append(('[외부사역]', False))
            right_lines += [(f'- {x}', False) for x in rep.external]
        if rep.internal:
            if right_lines: right_lines.append(('', False))
            right_lines.append(('[자체사역]', False))
            right_lines += [(f'- {x}', False) for x in rep.internal]
        _set_cell_lines(left, left_lines, bold_first=True)
        _set_cell_lines(right, right_lines, bold_first=True)

    # Weekly input contains no finance/prayer data: clear old template values, keep structure/headers.
    finance_idx = _finance_header_idx(table)
    prayer_idx = _prayer_row_idx(table)
    if finance_idx is not None:
        end = prayer_idx if prayer_idx is not None else len(table.rows)
        for idx in range(finance_idx + 1, end):
            row = table.rows[idx]
            uniques = _unique_tcs(row)
            for c in uniques[1:]:
                _clear_cell(c)
    if prayer_idx is not None:
        uniques = _unique_tcs(table.rows[prayer_idx])
        for c in uniques[1:]:
            _clear_cell(c)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _sundays(year: int, month: int) -> list[int]:
    cal = calendar.monthcalendar(year, month)
    # calendar module: Monday=0, Sunday=6
    return [week[calendar.SUNDAY] for week in cal if week[calendar.SUNDAY] != 0]


def _extract_week_cell_sections(text: str):
    section = None
    result = {'creator': [], 'whole': [], 'bulletin': {}, 'external': [], 'internal': [], 'other': []}
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith('<크리에이터>'):
            section = 'creator'; continue
        if line.startswith('<전체팀>') or line.startswith('< 전체팀>'):
            section = 'whole'; continue
        if line == '[주보]':
            section = 'bulletin'; continue
        if line == '[외부사역]':
            section = 'external'; continue
        if line == '[자체사역]':
            section = 'internal'; continue
        if line.startswith('[') and line.endswith(']'):
            section = 'other'; continue
        line = re.sub(r'^-\s*', '', line)
        if section == 'bulletin':
            m = re.match(r'(전면|후면|내지|단추가)\s*[:：]\s*(.+)', line)
            if m: result['bulletin'][m.group(1)] = m.group(2).strip()
        elif section in result and isinstance(result[section], list):
            result[section].append(line)
    return result


def parse_monthly_docx(data: bytes, year: int | None = None, month_hint: int | None = None) -> MonthlyData:
    doc = Document(io.BytesIO(data))
    table = doc.tables[0]
    title = ' '.join(c.text for c in _unique_tcs(table.rows[0]))
    mm = re.search(r'(\d{1,2})월', title)
    month = int(mm.group(1)) if mm else (month_hint or 1)
    year = year or date.today().year
    result = MonthlyData(year=year, month=month)
    sundays = _sundays(year, month)

    for row in table.rows:
        label = row.cells[0].text.strip()
        m = re.fullmatch(r'(\d+)주', label)
        if not m:
            continue
        week_no = int(m.group(1))
        uniques = _unique_tcs(row)
        combined = '\n'.join(c.text for c in uniques[1:])
        parts = _extract_week_cell_sections(combined)
        if parts['bulletin']:
            day = sundays[week_no - 1] if week_no <= len(sundays) else None
            result.bulletin_rows.append({'day': day, **parts['bulletin']})
        result.whole_team.extend(parts['whole'])
        result.internal.extend(parts['internal'])
        # 외부사역은 여러 주에 같은 항목이 반복될 수 있으므로
        # 월/분기 보고서에는 동일 문구를 최초 1회만 유지한다.
        # 의미가 비슷하더라도 문구가 다르면 임의로 합치지 않는다.
        for item in parts['external']:
            normalized = re.sub(r'\s+', ' ', item).strip()
            existing = {re.sub(r'\s+', ' ', x).strip() for x in result.external}
            if normalized and normalized not in existing:
                result.external.append(item.strip())

    # Finance total if present.
    for row in table.rows:
        if any(c.text.strip() == '합 계' for c in row.cells):
            nums = [re.sub(r'[^0-9,]', '', c.text) for c in row.cells]
            nums = [n for n in nums if n]
            if nums:
                result.finance_total = nums[-1]
    pidx = _prayer_row_idx(table)
    if pidx is not None:
        txt = '\n'.join(c.text for c in _unique_tcs(table.rows[pidx])[1:]).strip()
        if txt:
            result.prayers = [x.strip() for x in re.split(r'\n(?=\d+\.)', txt) if x.strip()]
    return result


def quarter_for_month(month: int) -> tuple[int, list[int]]:
    mapping = {1: (1, [12,1,2]), 2:(1,[12,1,2]), 12:(1,[12,1,2]),
               3:(2,[3,4,5]),4:(2,[3,4,5]),5:(2,[3,4,5]),
               6:(3,[6,7,8]),7:(3,[6,7,8]),8:(3,[6,7,8]),
               9:(4,[9,10,11]),10:(4,[9,10,11]),11:(4,[9,10,11])}
    return mapping[month]


def _set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def _add_text(cell, text: str, bold=False, size=9):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def build_quarterly_docx(month_docs: list[tuple[bytes, int, int]]) -> bytes:
    """month_docs: [(docx_bytes, year, month_hint), ...]. Generates DOCX in quarter-report structure."""
    months = [parse_monthly_docx(b, year=y, month_hint=m) for b, y, m in month_docs]
    months.sort(key=lambda x: x.month)
    # Determine quarter from first month; caller can upload in any order.
    q, qmonths = quarter_for_month(months[0].month)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Pt(35)
    sec.left_margin = sec.right_margin = Pt(36)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{q} / 4 확대기획위원회')
    r.bold = True; r.font.size = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('미디어팀')
    r.bold = True; r.font.size = Pt(16)

    for md in months:
        h = doc.add_paragraph()
        rr = h.add_run(f'{md.month}월')
        rr.bold = True; rr.font.size = Pt(12)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0]
        _set_repeat_table_header(hdr)
        for cell, text in zip(hdr.cells, ['날짜', '표지(전면)', '이미지광고(후면)', '단추가/간지/내지']):
            cell.text = ''
            _add_text(cell, text, bold=True)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for b in md.bulletin_rows:
            row = table.add_row().cells
            day = b.get('day')
            row[0].text = f'{md.month:02d}.{day:02d}' if day else ''
            row[1].text = b.get('전면', '')
            row[2].text = b.get('후면', '')
            inner = []
            if b.get('내지'): inner.append(f"내) {b['내지']}")
            if b.get('단추가'): inner.append(f"단) {b['단추가']}")
            row[3].text = '\n'.join(inner)
            for c in row:
                c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in c.paragraphs:
                    for run in p.runs: run.font.size = Pt(9)
        def add_section(title, items):
            if not items:
                return
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(f'▶{title}')
            r.bold = True; r.font.size = Pt(10)
            # exact content is retained, no deduplication that could lose source detail
            for item in items:
                pp = doc.add_paragraph(style=None)
                pp.paragraph_format.left_indent = Pt(10)
                pp.paragraph_format.space_after = Pt(0)
                rr = pp.add_run(item)
                rr.font.size = Pt(9)
        add_section('전체팀', md.whole_team)
        add_section('자체사역', md.internal)
        add_section('외부사역', md.external)

    # Finance: only values that actually exist in uploaded monthly reports.
    finance_months = [m for m in months if m.finance_total]
    if finance_months:
        p = doc.add_paragraph(); r = p.add_run('재정'); r.bold = True; r.font.size = Pt(12)
        t = doc.add_table(rows=1, cols=6); t.style = 'Table Grid'
        for c, txt in zip(t.rows[0].cells, ['월','이월금','수입','지출','반환','잔액']):
            c.text = txt
        for m in finance_months:
            cells = t.add_row().cells
            cells[0].text = f'{m.month}월'
            # Source monthly report only gives total reliably; do not invent other fields.
            cells[3].text = m.finance_total or ''

    prayer_source = next((m for m in reversed(months) if m.prayers), None)
    if prayer_source:
        p = doc.add_paragraph(); r = p.add_run('기도'); r.bold = True; r.font.size = Pt(12)
        for prayer in prayer_source.prayers:
            pp = doc.add_paragraph(prayer)
            for run in pp.runs: run.font.size = Pt(9)

    out = io.BytesIO(); doc.save(out); return out.getvalue()

# --- HWPX quarter output ----------------------------------------------------
import zipfile
from lxml import etree


def _local(el):
    return etree.QName(el).localname


def _direct_children(el, name):
    return [x for x in el if _local(x) == name]


def _text_of(el):
    return ''.join(el.xpath('.//*[local-name()="t"]/text()'))


def _first_t(el):
    ts = el.xpath('.//*[local-name()="t"]')
    return ts[0] if ts else None


def _set_text_keep_style(el, text: str):
    """Replace text while retaining the first run/paragraph formatting."""
    ts = el.xpath('.//*[local-name()="t"]')
    if not ts:
        return
    ts[0].text = text or ''
    for t in ts[1:]:
        t.text = ''
    # Cached line-layout information becomes stale after edits; Hancom recalculates it.
    for ls in el.xpath('.//*[local-name()="linesegarray"]'):
        parent = ls.getparent()
        if parent is not None:
            parent.remove(ls)


def _set_cell_paragraphs(cell, lines: list[str]):
    """Set one or more paragraphs in a HWPX table cell by cloning its existing style."""
    subs = _direct_children(cell, 'subList')
    if not subs:
        return
    sub = subs[0]
    paras = _direct_children(sub, 'p')
    if not paras:
        return
    proto = copy.deepcopy(paras[0])
    for p in paras:
        sub.remove(p)
    if not lines:
        lines = ['']
    insert_at = 0
    for line in lines:
        p = copy.deepcopy(proto)
        _set_text_keep_style(p, line)
        # keep paragraphs before any non-p children in subList
        sub.insert(insert_at, p)
        insert_at += 1


def _set_para_text(p, text: str):
    _set_text_keep_style(p, text)


def _nested_table(month_cell):
    tables = month_cell.xpath('.//*[local-name()="tbl"]')
    return tables[0] if tables else None


def _resize_hwpx_data_rows(tbl, wanted: int):
    """Keep header + wanted data rows, cloning the last styled row if needed."""
    rows = _direct_children(tbl, 'tr')
    if not rows:
        return
    current = len(rows) - 1
    while current < wanted:
        proto = copy.deepcopy(rows[-1])
        tbl.append(proto)
        rows = _direct_children(tbl, 'tr')
        current += 1
    while current > wanted:
        tbl.remove(rows[-1])
        rows = _direct_children(tbl, 'tr')
        current -= 1
    # normalize rowAddr values after resize
    rows = _direct_children(tbl, 'tr')
    for ri, row in enumerate(rows):
        cells = _direct_children(row, 'tc')
        for ci, cell in enumerate(cells):
            addrs = _direct_children(cell, 'cellAddr')
            if addrs:
                addrs[0].set('rowAddr', str(ri))
                addrs[0].set('colAddr', str(ci))


def _replace_month_section(month_cell, md: MonthlyData):
    """Replace section paragraphs after the nested bulletin table, preserving paragraph styles."""
    subs = _direct_children(month_cell, 'subList')
    if not subs:
        return
    sub = subs[0]
    paras = _direct_children(sub, 'p')
    # Paragraph 0 owns the nested table. Keep it and remove all later p nodes.
    owner = next((p for p in paras if p.xpath('.//*[local-name()="tbl"]')), paras[0] if paras else None)
    if owner is None:
        return

    # Pick styled prototypes from existing content.
    heading_proto = next((copy.deepcopy(p) for p in paras if _text_of(p).strip().startswith('▶')), copy.deepcopy(owner))
    content_proto = next((copy.deepcopy(p) for p in paras if _text_of(p).strip() and not _text_of(p).strip().startswith('▶') and not p.xpath('.//*[local-name()="tbl"]')), copy.deepcopy(owner))
    blank_proto = next((copy.deepcopy(p) for p in paras if not _text_of(p).strip() and not p.xpath('.//*[local-name()="tbl"]')), copy.deepcopy(content_proto))

    for p in list(_direct_children(sub, 'p')):
        if p is not owner:
            sub.remove(p)

    # Insert immediately after owner; avoid synthesizing empty sections.
    anchor_index = list(sub).index(owner) + 1
    new_paras = []
    def add_section(title, items, combine=False):
        nonlocal new_paras
        items = [x for x in items if x and x.strip()]
        if not items:
            return
        h = copy.deepcopy(heading_proto); _set_para_text(h, f'▶{title}'); new_paras.append(h)
        if combine:
            p = copy.deepcopy(content_proto); _set_para_text(p, ', '.join(items)); new_paras.append(p)
        else:
            for item in items:
                p = copy.deepcopy(content_proto); _set_para_text(p, item); new_paras.append(p)
    add_section('전체팀', md.whole_team)
    add_section('자체사역', md.internal, combine=True)
    add_section('외부사역', md.external, combine=True)
    b = copy.deepcopy(blank_proto); _set_para_text(b, ''); new_paras.append(b)
    for p in new_paras:
        sub.insert(anchor_index, p); anchor_index += 1


def _set_top_quarter(root, q: int):
    for tbl in root.xpath('//*[local-name()="tbl"]'):
        txt = _text_of(tbl)
        if '확대기획위원회' in txt:
            cells = tbl.xpath('./*[local-name()="tr"]/*[local-name()="tc"]')
            for cell in cells:
                if '확대기획위원회' in _text_of(cell):
                    _set_text_keep_style(cell, f'{q} / 4 확대기획위원회')
                    return


def _find_main_outer_table(root):
    for tbl in root.xpath('//*[local-name()="tbl"]'):
        rows = _direct_children(tbl, 'tr')
        if rows:
            first = ' '.join(_text_of(c) for c in _direct_children(rows[0], 'tc'))
            if '분기' in first and '사역내용' in first:
                return tbl
    raise ValueError('HWPX 템플릿에서 분기 사역내용 표를 찾지 못했습니다.')


def _finance_values_from_month_docx(data: bytes):
    """Extract explicitly present finance totals from the monthly DOCX sum row."""
    try:
        doc = Document(io.BytesIO(data))
    except Exception:
        return {}
    if not doc.tables:
        return {}
    table = doc.tables[0]
    header_row = None
    for row in table.rows:
        raw = [c.text.strip() for c in row.cells]
        if '구분' in raw and '수입' in raw and '지출' in raw:
            header_row = raw
            continue
        if header_row and any(v == '합 계' for v in raw):
            out = {}
            # Use physical grid positions, not unique merged cells. This keeps
            # the source column meaning intact even when Word cells are merged.
            for key in ('이월금', '수입', '지출', '잔액', '반환'):
                if key in header_row:
                    idx = header_row.index(key)
                    if idx < len(raw):
                        val = raw[idx].strip()
                        if val:
                            out[key] = val
            return out
    return {}


def build_quarterly_hwpx(template_path: str | Path, month_docs: list[tuple[bytes, int, int]]) -> bytes:
    """Create a quarter HWPX by editing only text/row count inside the user's HWPX template."""
    if len(month_docs) != 3:
        raise ValueError('분기 보고서는 월 보고서 3개가 필요합니다.')
    parsed = [(parse_monthly_docx(b, year=y, month_hint=m), b) for b, y, m in month_docs]
    qset = {quarter_for_month(md.month)[0] for md, _ in parsed}
    if len(qset) != 1:
        raise ValueError('세 월 보고서가 같은 분기에 속해야 합니다.')
    q = next(iter(qset))
    _, qmonths = quarter_for_month(parsed[0][0].month)
    order = {m:i for i,m in enumerate(qmonths)}
    parsed.sort(key=lambda pair: order[pair[0].month])

    with zipfile.ZipFile(str(template_path), 'r') as zin:
        files = {name: zin.read(name) for name in zin.namelist()}
    section_name = 'Contents/section0.xml'
    if section_name not in files:
        raise ValueError('HWPX 템플릿에 Contents/section0.xml이 없습니다.')
    parser = etree.XMLParser(remove_blank_text=False)
    root = etree.fromstring(files[section_name], parser)
    _set_top_quarter(root, q)
    outer = _find_main_outer_table(root)
    rows = _direct_children(outer, 'tr')
    if len(rows) < 8:
        raise ValueError('HWPX 분기 템플릿의 기본 행 구조가 예상과 다릅니다.')

    # First three rows after header are month blocks.
    month_rows = rows[1:4]
    for mr, (md, original_bytes) in zip(month_rows, parsed):
        cells = _direct_children(mr, 'tc')
        if len(cells) < 2:
            continue
        _set_cell_paragraphs(cells[0], [f'{md.month}월'])
        content_cell = cells[1]
        nested = _nested_table(content_cell)
        if nested is None:
            raise ValueError(f'{md.month}월 영역의 주보 표를 찾지 못했습니다.')
        sundays = _sundays(md.year, md.month)
        _resize_hwpx_data_rows(nested, len(sundays))
        by_day = {r.get('day'): r for r in md.bulletin_rows if r.get('day')}
        nrows = _direct_children(nested, 'tr')
        for i, day in enumerate(sundays, start=1):
            cells2 = _direct_children(nrows[i], 'tc')
            src = by_day.get(day, {})
            _set_cell_paragraphs(cells2[0], [f'{md.month:02d}.{day:02d}'])
            _set_cell_paragraphs(cells2[1], [src.get('전면','')])
            _set_cell_paragraphs(cells2[2], [src.get('후면','')])
            inner = []
            if src.get('내지'): inner.append(f"내) {src['내지']}")
            if src.get('단추가'): inner.append(f"단) {src['단추가']}")
            _set_cell_paragraphs(cells2[3], inner)
        _replace_month_section(content_cell, md)

    # Finance rows: month labels always update; unknown fields are cleared instead of invented.
    rows = _direct_children(outer, 'tr')
    finance_rows = rows[5:8]
    for fr, (md, original_bytes) in zip(finance_rows, parsed):
        cells = _direct_children(fr, 'tc')
        if not cells: continue
        _set_cell_paragraphs(cells[0], [f'{md.month}월'])
        values = _finance_values_from_month_docx(original_bytes)
        # Quarter template data columns are: month, carryover, income, expense, return, balance.
        keys = ['이월금','수입','지출','반환','잔액']
        for ci, key in enumerate(keys, start=1):
            if ci < len(cells):
                _set_cell_paragraphs(cells[ci], [values.get(key, '')])

    # Prayer block: use the latest uploaded month that explicitly contains prayers.
    rows = _direct_children(outer, 'tr')
    prayer_row = rows[8]
    pcells = _direct_children(prayer_row, 'tc')
    if len(pcells) >= 2:
        prayer_source = next((md for md, _ in reversed(parsed) if md.prayers), None)
        if prayer_source:
            _set_cell_paragraphs(pcells[1], prayer_source.prayers)
        else:
            _set_cell_paragraphs(pcells[1], [''])

    files[section_name] = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=None)
    out = io.BytesIO()
    with zipfile.ZipFile(out, 'w') as zout:
        # HWPX convention: mimetype first and uncompressed.
        if 'mimetype' in files:
            zout.writestr('mimetype', files['mimetype'], compress_type=zipfile.ZIP_STORED)
        for name, blob in files.items():
            if name == 'mimetype': continue
            zout.writestr(name, blob, compress_type=zipfile.ZIP_DEFLATED)
    return out.getvalue()
